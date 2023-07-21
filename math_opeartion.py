#addition
import random


data = []
group = []
while len(data) < 20:
    nums1 = random.randint(1, 9)
    nums2 = random.randint(1, 9)
    group.append(f"{nums1} + {nums2} = ")
    if len(group) == 4:
        data.append(group)
        group = []
print(data) 

#multiple
data_1 = []
group_1 = []
while len(data_1) < 20:
    nums3 = random.randint(1, 9)
    nums4 = random.randint(1, 9)
    group_1.append(f"{nums3} * {nums4} = ")
    if len(group_1) == 4:
        data_1.append(group_1)
        group_1 = []
print(data_1)

#export to word file
def word(data, file_name):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    for i in data:
        row = '\t\t'.join(i)
        doc.add_paragraph(row)
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = Pt(16)
    doc.save(f'{file_name}.docx')